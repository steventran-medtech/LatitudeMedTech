"""
Latitude MedTech — QMS Simulator Agent
=======================================
Generates a complete ISO 13485:2016 QMS document bundle from a plain-English
description of a medical device or pharmaceutical compound. Designed for
mock audit preparation and client QMS training.

Documents produced per run:
  00_DEVICE_PROFILE.json       — Device classification + all metadata
  DHR_batch_{lot}.docx         — Device History Record (primary deliverable)
  DMR_index.md                 — Device Master Record Index
  risk_management_summary.md   — ISO 14971 Risk Management File Summary
  CAPA_simulation.md           — 2–3 CAPA records for training scenarios
  sop_reference_list.md        — Applicable SOPs matrix
  training_records.md          — Personnel training qualification summary
  supplier_qualification.md    — Approved Supplier List extract
  QMS_BUNDLE_INDEX.md          — Master index of all generated files

Usage:
  python qms_simulator_agent.py "Insulin Delivery Pen"
  python qms_simulator_agent.py "Continuous Glucose Monitor" --out C:/custom/path
  python qms_simulator_agent.py "Titanium Knee Implant System" --profile-only

Interpreter: Use Athena/voice/venv (has anthropic, python-docx, python-dotenv)
  .\\Athena\\voice\\venv\\Scripts\\python.exe qms_simulator_agent.py "..."

DISCLAIMER: All generated content is fictional and for educational / mock audit
training purposes only. Not actual regulatory documentation. Not regulatory or
legal advice.  Label: Alpha — Steve Review Required.
"""

# ────────────────────────────────────────────────────────────────────────────
# Imports
# ────────────────────────────────────────────────────────────────────────────
import os
import sys
import json
import re
import argparse
import logging
from pathlib import Path
from datetime import datetime
from typing import Optional, List, Dict, Any

from dotenv import load_dotenv
import anthropic

# ── Path setup ───────────────────────────────────────────────────────────────
_AGENTS_DIR   = Path(__file__).resolve().parent
_ATHENA_ROOT  = _AGENTS_DIR.parent
_ENV_FILE     = _ATHENA_ROOT / "voice" / ".env"

load_dotenv(_ENV_FILE)
sys.path.insert(0, str(_AGENTS_DIR))

# Optional: integrate with Athena memory if available
try:
    from agent_base import AgentBase
    _BASE = AgentBase("qms_simulator")
except Exception:
    _BASE = None

try:
    from memory import Memory
    _MEM = Memory()
except Exception:
    _MEM = None

# ── Output / log directories ─────────────────────────────────────────────────
QMS_DIR = _ATHENA_ROOT / "documents" / "qms"
LOG_DIR = _ATHENA_ROOT / "logs"
QMS_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)

# ── Logging ──────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)s  %(message)s",
    handlers=[
        logging.FileHandler(LOG_DIR / "qms_simulator.log", encoding="utf-8"),
        logging.StreamHandler(sys.stdout),
    ]
)
log = logging.getLogger("qms_simulator")

# ── Models ───────────────────────────────────────────────────────────────────
MODEL_PROFILE  = "claude-sonnet-4-6"   # JSON profile — needs precision
MODEL_DOCUMENT = "claude-sonnet-4-6"   # Document content — needs quality

ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

# ── Branding / disclaimer ────────────────────────────────────────────────────
LATITUDE_NAVY = "1F3A6B"
DISCLAIMER_DOCX = (
    "SIMULATED DOCUMENT — MOCK AUDIT TRAINING USE ONLY  |  "
    "Latitude MedTech LLC · Educational / Planning Purposes · Not Regulatory Documentation  |  "
    "Alpha — Steve Review Required"
)
DISCLAIMER_MD = (
    "\n\n---\n\n"
    "_**DISCLAIMER:** Produced by Latitude MedTech LLC for educational and mock audit training "
    "purposes only. All company names, personnel names, lot numbers, device data, and "
    "document numbers are entirely fictional. This is NOT actual regulatory documentation "
    "and does NOT constitute regulatory or legal advice._\n\n"
    "_Label: Alpha — Steve Review Required._"
)

# ── QMS system prompts ────────────────────────────────────────────────────────
# Used for all narrative document generation
QMS_SYSTEM = """You are a QMS documentation specialist at Latitude MedTech LLC producing
simulated ISO 13485:2016 / 21 CFR Part 820 quality records for mock audit training.

Rules:
- Use authentic regulatory language, realistic device-specific technical data, and
  plausible manufacturing parameters. Outputs must read like real QMS documents.
- All company names, personnel names, lot numbers, clearance numbers, and values must
  be entirely fictional. Never use real company names or real FDA clearance numbers.
- Documents must be internally consistent — use the exact names, numbers, and
  identifiers from the provided device profile.
- Output ONLY the requested content — no preamble, no meta-commentary."""

# Used ONLY for JSON profile generation — no document prefix, pure JSON output
PROFILE_SYSTEM = """You are a medical device classification and QMS metadata specialist.
You generate structured JSON data for QMS simulation systems.

Rules:
- Output ONLY a valid JSON object. Absolutely no text before or after the JSON.
- No markdown code fences. No preamble. No explanations. Just the raw JSON object.
- All company names, personnel names, lot numbers, and clearance numbers must be fictional.
  Never use real MedTech company names or real FDA clearance numbers.
- Be technically accurate for the stated device type — use realistic product codes,
  applicable standards, manufacturing steps, and regulatory classifications."""


# ════════════════════════════════════════════════════════════════════════════
# PROMPT TEMPLATES
# ════════════════════════════════════════════════════════════════════════════

PROFILE_PROMPT = """Generate a comprehensive, internally consistent device profile JSON for a
QMS simulation. This profile will drive generation of all QMS documents for a mock audit.

DEVICE DESCRIPTION: {description}

Return ONLY a valid JSON object (no markdown fences, no preamble) matching this schema:

{{
  "device_name": "generic device name",
  "trade_name": "fictional trade/brand name",
  "model_number": "part/model number",
  "manufacturer_name": "fictional company name",
  "manufacturer_address": "fictional US street address, City, ST ZIP",
  "fda_establishment_reg": "10-digit fictional FDA establishment reg number",
  "device_class": "FDA Class I / II / III  OR  Rx Pharmaceutical / OTC Pharmaceutical",
  "regulatory_pathway": "510(k) / PMA / De Novo / NDA / ANDA / 510(k) Exempt",
  "clearance_number": "fictional clearance number e.g. K252345 or NDA-123456",
  "product_code": "3-letter FDA product code realistic for this device type",
  "regulation_number": "applicable 21 CFR part/section e.g. 21 CFR 880.5860",
  "is_sterile": true_or_false,
  "is_implant": true_or_false,
  "has_software": true_or_false,
  "is_combination_product": true_or_false,
  "is_pharma": true_or_false,
  "intended_use": "2-3 sentence intended use statement",
  "indications_for_use": "2-3 sentence indications for use",
  "applicable_standards": ["ISO 13485:2016", "at least 3-5 device-specific standards"],
  "mock_lot_number": "LOT-2025-XXXX format",
  "mock_serial_prefix": "short alphanumeric prefix",
  "mock_batch_quantity": realistic_integer,
  "mock_manufacture_date": "2025-10-15",
  "mock_expiration_date": "2027-10-14 or N/A if no expiry",
  "personnel": {{
    "qa_manager": "fictional full name",
    "production_supervisor": "fictional full name",
    "quality_engineer": "fictional full name",
    "regulatory_affairs": "fictional full name",
    "design_engineer": "fictional full name or null"
  }},
  "critical_components": [
    {{
      "component": "component description",
      "part_number": "P/N format",
      "revision": "Rev. X",
      "supplier": "fictional supplier name",
      "supplier_item_number": "supplier's item number",
      "mock_lot": "supplier lot number",
      "material": "material or substance",
      "is_critical": true_or_false,
      "coa_required": true_or_false
    }}
  ],
  "manufacturing_steps": [
    {{
      "step": 1,
      "operation": "operation description",
      "work_instruction": "WI-XXX",
      "equipment": "equipment name + model",
      "equipment_id": "EQ-XXX",
      "acceptance_criteria": "measurable pass/fail criteria",
      "inspection_type": "First Article / In-Process / Final",
      "mock_result": "specific measured value or observation",
      "status": "PASS",
      "operator": "fictional first + last initial or name",
      "date": "2025-10-15"
    }}
  ],
  "in_process_inspections": [
    {{
      "inspection_point": "inspection point name",
      "specification": "spec with units and tolerances",
      "method": "inspection method",
      "actual_result": "measured value",
      "status": "PASS",
      "inspector": "fictional name",
      "date": "2025-10-15"
    }}
  ],
  "applicable_sops": [
    {{
      "number": "SOP-XXX",
      "title": "SOP title",
      "version": "Rev. X.X",
      "department": "QA / Production / Engineering / RA"
    }}
  ],
  "dmr_documents": [
    {{
      "document_number": "DMR-XXX",
      "title": "document title",
      "revision": "Rev. X",
      "type": "Specification / Drawing / Work Instruction / Process / Validation"
    }}
  ],
  "top_hazards": [
    {{
      "hazard_id": "H-001",
      "hazard": "hazard source",
      "hazardous_situation": "hazardous situation description",
      "harm": "harm to patient or user",
      "severity": "Catastrophic / Critical / Serious / Minor / Negligible",
      "probability": "Frequent / Probable / Occasional / Remote / Improbable",
      "initial_risk": "Unacceptable / ALARP / Acceptable",
      "risk_control": "risk control measure applied",
      "residual_risk": "ALARP / Acceptable"
    }}
  ],
  "suppliers": [
    {{
      "supplier_id": "SUP-XXX",
      "name": "fictional supplier name",
      "items_supplied": "description of items",
      "criticality": "Critical / Major / Minor",
      "qualification_status": "Approved / Conditional / Under Review",
      "last_audit_date": "2024-11-15",
      "re_qual_due": "2026-11-15",
      "iso_13485_certified": true_or_false
    }}
  ],
  "capa_scenarios": [
    {{
      "capa_number": "CAPA-2025-XXX",
      "source": "Internal Audit / Customer Complaint / Supplier NC / Management Review",
      "title": "brief CAPA title",
      "description": "specific nonconformance description",
      "root_cause": "true root cause (not just symptom)",
      "corrective_action": "specific corrective action taken",
      "preventive_action": "specific systemic preventive action",
      "verification_method": "how effectiveness will be verified",
      "status": "Open / Closed",
      "opened_date": "2025-06-15",
      "closed_date": "2025-09-30 or null if Open"
    }}
  ]
}}

Generate exactly: 6 manufacturing_steps, 5 critical_components, 8 applicable_sops,
6 dmr_documents, 4 top_hazards, 4 suppliers, 2 capa_scenarios.
Keep string values concise — 1–2 sentences max for descriptions.
Be specific and technically accurate for this device/compound type.
Use plausible fictional company and personnel names — do NOT use real MedTech companies."""


DMR_PROMPT = """Using the device profile JSON below, generate a Device Master Record (DMR) Index
per ISO 13485:2016 §4.2.3 and 21 CFR 820.181.

DEVICE PROFILE:
{profile_json}

---

Begin with: **SIMULATED RECORD — MOCK AUDIT USE ONLY**

Then generate the full DMR Index in Markdown:

## Device Master Record Index
**Document Number:** DMR-IDX-{model_number} | **Revision:** Rev. 1.0 | **Effective Date:** {manufacture_date}
**Prepared By:** {qa_manager} | **Approved By:** {qa_manager} | **Status:** Released

---

### Section 1 — Device Identification
(Device name, trade name, model/part number, FDA product code, classification, regulatory pathway,
applicable standards — pulled directly from the profile)

### Section 2 — DMR Constituent Documents
A properly formatted table with ALL documents from the profile's dmr_documents list plus any
additional documents logically required for this device type.

| Document Number | Title | Revision | Type | Owner | Location |
|---|---|---|---|---|---|

Include at minimum (add to profile list if missing):
- Device specification / product standard
- Engineering drawing or dimensional spec
- Master bill of materials
- All manufacturing work instructions
- Incoming inspection procedure
- Final acceptance test procedure
- Packaging and labeling specification
- Sterilization specification (if is_sterile = true, else mark N/A)
- Software documentation (if has_software = true, else mark N/A)

### Section 3 — Applicable Standards and Regulations
| Standard / Regulation | Applicable Clauses | Compliance Method |
|---|---|---|

### Section 4 — Associated Records
Types of DHRs maintained, their retention period, Design History File location (if applicable),
installation qualification records (if applicable).

### Section 5 — Document Control Summary
How DMR documents are revised, who approves changes, where the master copy is stored,
distribution controls.

Total: 600–900 words. Use exact document numbers from the profile."""

RISK_PROMPT = """Using the device profile JSON below, generate a Risk Management File Summary
per ISO 14971:2019.

DEVICE PROFILE:
{profile_json}

---

Begin with: **SIMULATED RECORD — MOCK AUDIT USE ONLY**

Then generate the Risk Management Summary in Markdown:

## Risk Management File Summary
**Document Number:** RMF-{model_number} | **Revision:** Rev. 1.0 | **Effective Date:** {manufacture_date}
**Prepared By:** {quality_engineer} | **Reviewed By:** {qa_manager}

---

### Section 1 — Risk Management Plan
Scope of the risk management process, intended use and foreseeable misuse, lifecycle stages
covered, risk acceptability criteria (define 3x3 or 5x5 risk matrix with thresholds), references
to risk management SOP.

### Section 2 — Risk Analysis Summary
| Hazard ID | Hazard | Hazardous Situation | Harm | Severity | Probability | Initial Risk |
|---|---|---|---|---|---|---|

Use ALL hazards from top_hazards in the profile. Add 1–2 more hazards specific to this device type.

### Section 3 — Risk Evaluation and Risk Controls
| Hazard ID | Initial Risk | Risk Control Measure | Control Type | Residual Risk | Acceptable? |
|---|---|---|---|---|---|

Control types: (1) Inherent safety by design, (2) Protective measures, (3) Information for safety

### Section 4 — Benefit-Risk Analysis
One paragraph assessing whether the clinical benefits outweigh the residual risks for this device.
Be specific to the device's patient population and therapeutic purpose.

### Section 5 — Overall Residual Risk Evaluation
Statement that after application of all risk controls, the overall residual risk is acceptable.
Reference ISO 14971 §8 and any post-market surveillance plan.

### Section 6 — Risk Management Review Statement
Formal sign-off statement: "The risk management file for [device name] has been reviewed.
Overall residual risk is [accepted/ALARP]. Signed: [QA Manager name], [date]."

Total: 700–1000 words. Use exact hazard IDs from the profile."""

CAPA_PROMPT = """Using the device profile JSON below, generate 2–3 detailed CAPA records
per ISO 13485:2016 §8.5.2/8.5.3 for mock audit training.

DEVICE PROFILE:
{profile_json}

---

Begin with: **SIMULATED RECORDS — MOCK AUDIT USE ONLY**

For EACH CAPA scenario in the profile, generate a complete CAPA record:

---

## [CAPA Number]: [CAPA Title]

| Field | Value |
|---|---|
| CAPA Number | [from profile] |
| Date Opened | [from profile] |
| Source | [Internal Audit / Customer Complaint / etc.] |
| Priority | High / Medium / Low (assess based on risk) |
| Status | Open / Closed |
| CAPA Owner | [relevant person from personnel] |
| Related NCR / Complaint | NCR-2025-XXX or COMP-2025-XXX |

### 1. Problem Description
[Specific, factual description — what was observed, where, when, how many units/occurrences]

### 2. Immediate Containment Action
[What was done within 24–72 hours to contain the problem and protect patients/users]

### 3. Root Cause Analysis

**Method:** [5-Why / Ishikawa (Fishbone) / Fault Tree — pick the appropriate method]

[Walk through the analysis step by step. Go to the true root cause, not just the symptom.
For 5-Why, show all 5 levels. For fishbone, list causes by category.]

**Root Cause Statement:** [One clear sentence: "The root cause is..."]

### 4. Corrective Action Plan
| # | Action Item | Responsible | Target Date | Completion Date | Evidence |
|---|---|---|---|---|---|
| CA-1 | [specific action] | [name] | [date] | [date or TBD] | [document/record] |
| CA-2 | [specific action] | [name] | [date] | [date or TBD] | [document/record] |

### 5. Preventive Action Plan
| # | Systemic Prevention | Responsible | Target Date | Completion Date |
|---|---|---|---|---|
| PA-1 | [systemic change to prevent recurrence] | [name] | [date] | [date or TBD] |

### 6. Effectiveness Verification
**Verification Method:** [specific — e.g., review 30 consecutive DHRs, audit 3 training records]
**Verification Criteria:** [what constitutes "effective"?]
**Verification Period:** [e.g., 90 days post-implementation]
**Verification Result:** [if Closed: what was found and conclusion; if Open: "Pending"]

### 7. CAPA Closure
| Field | Value |
|---|---|
| Verified Effective? | Yes / No / Pending |
| Closure Date | [from profile or "Open"] |
| Closed By | [QA Manager if closed, or "Pending"] |

---

Make each CAPA realistic and instructive. Each should illustrate a different source
(Internal Audit, Customer Complaint, Supplier NC). Root cause analysis must be specific —
not generic ISO platitudes. Total: 900–1300 words."""

SOP_PROMPT = """Using the device profile JSON below, generate an SOP Reference Matrix and
procedure summaries per ISO 13485:2016 §4.2.

DEVICE PROFILE:
{profile_json}

---

Begin with: **SIMULATED RECORD — MOCK AUDIT USE ONLY**

## Applicable Standard Operating Procedures — {device_name}
**Document Number:** SOP-MTX-{model_number} | **Revision:** Rev. 1.0 | **Effective Date:** {manufacture_date}
**Prepared By:** {qa_manager}

---

### Section 1 — SOP Matrix
A complete table of ALL SOPs from the profile's applicable_sops list.

| SOP Number | Title | Version | Department | Scope | Training Required? | Last Review Date |
|---|---|---|---|---|---|---|

Add any additional mandatory ISO 13485 SOPs not in the profile list:
- Document and Record Control
- Management Review
- Internal Audit
- CAPA
- Complaint Handling
- Risk Management
- Supplier Evaluation and Control

### Section 2 — Key SOP Summaries for This Device

For each of these mandatory SOPs (using the actual SOP number from the matrix):
**[SOP Number] — [Title]**
2–3 sentences explaining the key requirements as they specifically apply to manufacturing
and quality control of this device type.

Cover: Document Control, CAPA, Supplier Qualification, Internal Audit, Complaint Handling,
and the most important device-type-specific production SOP.

### Section 3 — SOP Change History
| SOP Number | Change Description | Effective Date | Changed By |
|---|---|---|---|
(Show 2–3 recent fictional SOP revisions relevant to this device)

Total: 550–750 words."""

TRAINING_PROMPT = """Using the device profile JSON below, generate Personnel Training Records
per ISO 13485:2016 §6.2.

DEVICE PROFILE:
{profile_json}

---

Begin with: **SIMULATED RECORD — MOCK AUDIT USE ONLY**

## Training Records Summary — {device_name} Production
**Document Number:** TRN-{model_number} | **Effective Date:** {manufacture_date}
**Prepared By:** {qa_manager}

---

### Section 1 — Personnel Qualification Matrix
A training completion table for ALL personnel in the profile.

| Employee Name | Title | SOP / WI Number | Training Method | Date Completed | Trainer | Pass/Fail | Retraining Due |
|---|---|---|---|---|---|---|---|

Include:
- QA Manager: trained on all quality/management SOPs
- Production Supervisor: trained on all production WIs + quality SOPs
- Quality Engineer: trained on inspection SOPs and audit procedures
- 2 fictional production operators (name them): trained on production WIs
- Regulatory Affairs: trained on regulatory SOPs

Training methods: Classroom, OJT (On-the-Job Training), Read & Understand, Practical Assessment

### Section 2 — Competency Assessment Records
For 2 critical operations specific to this device, show a brief competency assessment:

**Operation:** [step from manufacturing_steps]
**Trainee:** [name]
**Assessor:** [name]
**Method:** Practical demonstration
**Criteria:** [acceptance criteria from manufacturing_steps]
**Result:** Competent / Not Competent
**Comments:** [specific observation]

### Section 3 — Training Gap Analysis
One paragraph: upcoming retraining or new training required (new SOP revision, new
equipment, regulatory change). Be specific to this device type.

Total: 500–700 words."""

SUPPLIER_PROMPT = """Using the device profile JSON below, generate Supplier Qualification Records
per ISO 13485:2016 §7.4.

DEVICE PROFILE:
{profile_json}

---

Begin with: **SIMULATED RECORD — MOCK AUDIT USE ONLY**

## Supplier Qualification Records — {device_name}
**Document Number:** SQR-{model_number} | **Effective Date:** {manufacture_date}
**Prepared By:** {quality_engineer} | **Approved By:** {qa_manager}

---

### Section 1 — Approved Supplier List (ASL) — Extract
A table for ALL suppliers in the profile plus any additional suppliers implied by
the critical_components list.

| Supplier ID | Company Name | Items Supplied | Criticality | Qualification Method | Qual Date | Last Audit | Re-qual Due | Status | ISO 13485? |
|---|---|---|---|---|---|---|---|---|---|

### Section 2 — Critical Supplier Profiles
For each CRITICAL supplier (top 2–3 from the profile):

**Supplier: [Name]** (Supplier ID: SUP-XXX)
- **Items Supplied:** [list]
- **Qualification Basis:** [what documents / audits qualified them]
- **Performance — Last 12 Months:**
  - On-Time Delivery: XX%
  - Lot Acceptance Rate: XX%
  - Open NCRs: X
- **Last Audit Date:** [date] | **Audit Type:** [On-site / Remote / Desk]
- **Key Audit Findings:** [1–2 minor observations and their resolution]
- **Next Scheduled Audit:** [date]

### Section 3 — Supplier Monitoring Program
2–3 sentences: how suppliers are monitored on an ongoing basis (CoA review, incoming
inspection frequency, periodic audit cadence, performance scorecard).

### Section 4 — Supplier Nonconformance Records
1–2 fictional supplier NCs raised in the past 12 months with resolution.

| NCR Number | Supplier | Description | Date Issued | Root Cause | Resolution | Status |
|---|---|---|---|---|---|---|

Total: 650–850 words."""


TRACEABILITY_PROMPT = """Using the device profile JSON below, generate a Design and Risk Control
Traceability Matrix per ISO 13485:2016 §7.3 and ISO 14971:2019.

DEVICE PROFILE:
{profile_json}

---

Begin with: **SIMULATED RECORD — MOCK AUDIT USE ONLY**

## Design & Risk Control Traceability Matrix
**Document Number:** TRM-{model_number} | **Revision:** Rev. 1.0 | **Effective Date:** {manufacture_date}
**Prepared By:** {quality_engineer} | **Approved By:** {qa_manager}

---

### Section 1 — Design Input to Output Traceability
A table linking device design inputs to the design outputs that satisfy them,
the verification activity performed, and verification result.

| Req. ID | Design Input | Source | Design Output (DMR Doc #) | Verification Activity | Verification Result | Status |
|---|---|---|---|---|---|---|

Generate 10–14 requirements covering:
- User needs (patient/clinician usability, ergonomics)
- Functional performance (accuracy, dose precision, force, durability)
- Safety requirements — one row per hazard risk control from top_hazards
- Regulatory and standards requirements (one row per applicable_standard)
- Interface requirements (compatibility with accessories or consumables)
- Labeling and IFU requirements
- Packaging and environmental requirements

Use exact DMR document numbers from the profile's dmr_documents list.

### Section 2 — Risk Control Traceability
Links each risk control in the Risk Management File to the design output that
implements it and the verification evidence that it works.

| Hazard ID | Harm | Risk Control Requirement | Implementing Design Output (DMR Doc #) | Verification Method | Verification Document | Residual Risk | Verified? |
|---|---|---|---|---|---|---|---|

One row per hazard in the profile's top_hazards.

### Section 3 — Standards and Regulatory Requirements Traceability
| Standard / Regulation | Clause | Requirement Summary | Evidence / Design Output | Compliant? |
|---|---|---|---|---|

Cover 8–10 key clauses from the applicable_standards. Focus on safety,
performance, and labeling requirements — not an exhaustive clause list.

### Section 4 — Design Change Impact Assessment
Show 1–2 design changes consistent with the CAPA scenarios in the profile
(reference the ECO numbers or corrective actions from the CAPAs).

| Change Ref. | Change Description | Affected Design Inputs (Req. ID) | Reverification Required? | Reverification Status |
|---|---|---|---|---|

### Section 5 — Traceability Completeness Statement
Formal statement that all design inputs have been traced to verified design outputs
and that all risk controls have been verified as implemented. Signed by QA Manager.

Total: 800–1100 words. Use exact document numbers and hazard IDs from the profile."""


RISK_ANALYSIS_PROMPT = """Using the device profile JSON below, generate a Risk Analysis document
combining a Hazard Analysis and Design FMEA per ISO 14971:2019 §5.

DEVICE PROFILE:
{profile_json}

---

Begin with: **SIMULATED RECORD — MOCK AUDIT USE ONLY**

## Risk Analysis — Hazard Analysis and Design FMEA
**Document Number:** RA-{model_number} | **Revision:** Rev. 1.0 | **Effective Date:** {manufacture_date}
**Prepared By:** {quality_engineer} | **Reviewed By:** {qa_manager}

---

### Section 1 — Scope and Methodology
Device description, lifecycle stages covered, analysis team, standards referenced.

**Risk Estimation Matrix (5x5):**
Severity rows (1 Negligible → 5 Catastrophic) × Probability columns (1 Improbable → 5 Frequent).
Show threshold labels: scores 1–3 = Acceptable, 4–8 = ALARP, 9–25 = Unacceptable.

### Section 2 — Hazard Analysis
All hazards from the profile's top_hazards PLUS 3–4 additional device-specific hazards
(use errors, environmental hazards, maintenance hazards, end-of-life). Total 7–10 rows.

| Hazard ID | Hazard Source | Hazardous Situation | Foreseeable Harm | Severity (S) | Probability Before Controls (P) | Risk Score | Risk Level |
|---|---|---|---|---|---|---|---|

### Section 3 — Design FMEA
For 8–10 key device functions (drawn from manufacturing_steps and intended use),
analyze potential failure modes. Include RPN = Severity × Occurrence × Detection (each 1–5).
RPN ≥ 50 requires a risk action.

| FMEA ID | Component / Function | Failure Mode | Effect on Patient/User | S (1–5) | Cause | Occurrence (O 1–5) | Detection Method | Detection (D 1–5) | RPN | Action Required? | Risk Control |
|---|---|---|---|---|---|---|---|---|---|---|---|

### Section 4 — Manufacturing Process FMEA Extract
For the 3 highest-risk manufacturing steps in the profile, assess process failure modes.

| Step # | Process Operation | Failure Mode | Effect | S | Cause | O | Current Controls | D | RPN | Action? |
|---|---|---|---|---|---|---|---|---|---|---|

### Section 5 — Post-Control Risk Summary
For each hazard in Section 2, show residual risk after all risk controls are applied.

| Hazard ID | Initial Risk Level | Risk Control Applied | Post-Control S | Post-Control P | Residual Risk Score | Residual Level | Acceptable? | Reference |
|---|---|---|---|---|---|---|---|---|

### Section 6 — Risk Analysis Completeness Statement
Statement that the analysis is complete for the current design. Note any assumptions,
boundaries, or deferred analyses. Signed by QA Manager.

Total: 900–1200 words. Use exact hazard IDs from the profile's top_hazards."""


NCMR_TRENDING_PROMPT = """Using the device profile JSON below, generate an NCMR and Complaint
Trending Report per ISO 13485:2016 §8.3, §8.2.1, §8.4 and 21 CFR 820.198/820.90.

DEVICE PROFILE:
{profile_json}

---

Begin with: **SIMULATED RECORD — MOCK AUDIT USE ONLY**

## NCMR and Complaint Trending Report
**Document Number:** TRD-{model_number} | **Reporting Period:** January 1 – December 31, 2025
**Prepared By:** {quality_engineer} | **Reviewed By:** {qa_manager} | **Date:** {manufacture_date}

---

### Section 1 — Nonconformance Material Reports (NCMRs)

#### 1.1 NCMR Log — 2025
Generate 6–8 realistic NCMRs for this device type spread across the year.

| NCMR # | Date | Lot/Batch | Description | Category | Qty Affected | Disposition | Root Cause Category | CAPA Raised? | Status |
|---|---|---|---|---|---|---|---|---|---|

Categories: Component / Assembly / Inspection / Labeling / Documentation / Process
Dispositions: Use As Is / Rework / Scrap / Return to Supplier
Root Cause Categories: Incoming Material / Process / Operator Error / Equipment / Design

At least one NCMR should reference the CAPA scenario from the profile.

#### 1.2 NCMR Monthly Trend
| Month | NCMRs Opened | NCMRs Closed | Category Breakdown |
|---|---|---|---|

#### 1.3 Pareto Analysis
Identify the top 2–3 NCMR categories by frequency, state their share of total NCMRs,
and assess whether the current rate warrants escalation to CAPA or management review.

---

### Section 2 — Complaint Log and Analysis

#### 2.1 Complaint Log — 2025
Generate 4–6 realistic complaints for this device type.

| Complaint # | Date Received | Reporter | Lot # | Description | Category | Injury/Malfunction? | MDR Reportable? | Investigation Status |
|---|---|---|---|---|---|---|---|---|

Reporter: Patient / Clinician / Distributor
Categories: Performance / Malfunction / Labeling / Packaging / Use Error
At least one complaint should be linked to the CAPA customer complaint scenario.

#### 2.2 MDR / Vigilance Assessment
For each complaint where MDR Reportable = Yes (include at least 1):
- MDR decision rationale (21 CFR 803 section applied)
- Timeline: date received → investigation complete → MDR filed
- Summary of device failure investigation and conclusion

#### 2.3 Complaint Trending
Total complaints by category, complaint rate per 1,000 units (use mock_batch_quantity),
and signal detection: were any thresholds exceeded?

---

### Section 3 — Quality Signal Analysis and Management Review Input

#### 3.1 Signal Detection
Apply these thresholds to the data above:
- Same failure mode in 3 or more complaints or NCMRs → CAPA trigger
- Any confirmed patient harm → immediate CAPA
- Negative trend over 2+ consecutive months → management review flag

State clearly which signals were triggered (if any) and what action was taken.

#### 3.2 Management Review KPI Summary
| KPI | 2025 Value | Prior Year | Trend | Action Required? |
|---|---|---|---|---|
| Total NCMRs | | | | |
| Total Complaints | | | | |
| MDRs Filed | | | | |
| CAPAs Opened | | | | |
| Open NCMRs at Year-End | | | | |
| Complaint Rate (per 1,000 units) | | | | |

Fill in realistic values consistent with the NCMRs and complaints generated above.
Prior year values should show a plausible comparison (slight improvement or stable).

#### 3.3 Key Findings and Recommended Actions
3 bullet points summarizing the most important quality findings from this period,
and the recommended actions for management review.

#### 3.4 Report Approval
Formal sign-off by QA Manager, Production Supervisor, and Regulatory Affairs.

Total: 950–1250 words."""


# ════════════════════════════════════════════════════════════════════════════
# DOCX BUILDER UTILITIES
# ════════════════════════════════════════════════════════════════════════════

def _docx_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def _add_table_borders(table):
    """Apply single-line borders to all table cells."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl    = table._tbl
    tblPr  = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _shade_cell(cell, fill_hex: str = LATITUDE_NAVY):
    """Fill a table cell with a background color."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _header_row(table, headers, col_widths_pct=None):
    """Write a navy-on-white header row with bold text."""
    from docx.shared import Pt, RGBColor

    row = table.rows[0]
    for i, text in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell)


def _data_row(table, values, row_idx, shade_even=True):
    """Write a data row, optionally shading even rows light grey."""
    from docx.shared import Pt

    row = table.add_row()
    fill = "F2F2F2" if (shade_even and row_idx % 2 == 0) else "FFFFFF"
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(8.5)
        if fill == "F2F2F2":
            _shade_cell(cell, "F2F2F2")
    return row


def _add_section_heading(doc, text: str, level: int = 1):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6B)
        run.font.size = Pt(13 if level == 1 else 11)


def _add_controlled_doc_header(doc, profile: dict, doc_number: str, doc_title: str):
    """Insert a standard controlled-document info block at the top of a document."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(f"SIMULATED RECORD — MOCK AUDIT USE ONLY")
    tr.bold = True
    tr.font.size = Pt(10)

    doc.add_paragraph()
    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = title2.add_run(doc_title.upper())
    r2.bold = True
    r2.font.size = Pt(16)

    doc.add_paragraph()

    # Metadata table (2-column)
    meta = [
        ("Manufacturer:",         profile.get("manufacturer_name", "")),
        ("Device Name:",          profile.get("device_name", "")),
        ("Trade Name / Model:",   f"{profile.get('trade_name','')} / {profile.get('model_number','')}"),
        ("Lot / Batch Number:",   profile.get("mock_lot_number", "")),
        ("Manufacture Date:",     profile.get("mock_manufacture_date", "")),
        ("Expiration Date:",      profile.get("mock_expiration_date", "")),
        ("Batch Quantity:",       str(profile.get("mock_batch_quantity", ""))),
        ("Sterility Status:",     "Sterile" if profile.get("is_sterile") else "Non-Sterile"),
        ("Document Number:",      doc_number),
        ("Revision:",             "Rev. 1.0"),
        ("Effective Date:",       profile.get("mock_manufacture_date", "")),
        ("Prepared By:",          profile.get("personnel", {}).get("quality_engineer", "")),
        ("Reviewed By:",          profile.get("personnel", {}).get("production_supervisor", "")),
        ("Approved By:",          profile.get("personnel", {}).get("qa_manager", "")),
    ]
    tbl = doc.add_table(rows=1, cols=4)
    _add_table_borders(tbl)
    hdr_row = tbl.rows[0]
    for i, lbl in enumerate(["Field", "Value", "Field", "Value"]):
        hdr_row.cells[i].text = lbl
        hdr_row.cells[i].paragraphs[0].runs[0].bold = True
        _shade_cell(hdr_row.cells[i])

    left  = meta[:len(meta)//2 + len(meta)%2]
    right = meta[len(meta)//2 + len(meta)%2:]
    for idx in range(max(len(left), len(right))):
        row = tbl.add_row()
        lk, lv = left[idx]  if idx < len(left)  else ("", "")
        rk, rv = right[idx] if idx < len(right) else ("", "")
        row.cells[0].text = lk;  row.cells[0].paragraphs[0].runs[0].bold = True if lk else False
        row.cells[1].text = lv
        row.cells[2].text = rk;  row.cells[2].paragraphs[0].runs[0].bold = True if rk else False
        row.cells[3].text = rv
        if idx % 2 == 0:
            for c in row.cells:
                _shade_cell(c, "F2F2F2")

    doc.add_paragraph()


def _add_disclaimer_footer(doc):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_paragraph()
    hr = doc.add_paragraph("─" * 80)
    hr.alignment = WD_ALIGN_PARAGRAPH.CENTER

    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = disc.add_run(DISCLAIMER_DOCX)
    dr.font.size = Pt(8)
    dr.italic = True


# ════════════════════════════════════════════════════════════════════════════
# DHR DOCX BUILDER
# ════════════════════════════════════════════════════════════════════════════

def build_dhr_docx(profile: dict, out_path: Path):
    """
    Build a Device History Record as a professional Word document from the device profile.
    The profile's manufacturing_steps, critical_components, and in_process_inspections
    drive the content — no additional Claude call needed.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(0.75)

    # ── Cover / header ───────────────────────────────────────────────────────
    lot = profile.get("mock_lot_number", "LOT-XXXX")
    _add_controlled_doc_header(
        doc, profile,
        doc_number=f"DHR-{lot}",
        doc_title="Device History Record (DHR)"
    )

    # ── Section 1: Manufacturing Traveler ─────────────────────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "Section 1 — Manufacturing Traveler")

    steps = profile.get("manufacturing_steps", [])
    if steps:
        cols = ["Step", "Operation", "Work Instruction", "Equipment (ID)", "Acceptance Criteria",
                "Actual Result", "Status", "Operator", "Date"]
        tbl = doc.add_table(rows=1, cols=len(cols))
        _add_table_borders(tbl)
        _header_row(tbl, cols)

        for idx, step in enumerate(steps):
            _data_row(tbl, [
                str(step.get("step", idx + 1)),
                step.get("operation", ""),
                step.get("work_instruction", ""),
                f"{step.get('equipment','')} ({step.get('equipment_id','')})",
                step.get("acceptance_criteria", ""),
                step.get("mock_result", ""),
                step.get("status", "PASS"),
                step.get("operator", ""),
                step.get("date", ""),
            ], idx)
    else:
        doc.add_paragraph("No manufacturing steps defined in device profile.")

    # ── Section 2: Bill of Materials — Lot Traceability ──────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "Section 2 — Bill of Materials: Lot Traceability")

    components = profile.get("critical_components", [])
    if components:
        cols = ["Component", "Part Number / Rev.", "Supplier", "Supplier Item #",
                "Supplier Lot #", "Qty / Unit", "Material", "CoA Required?", "CoA Reviewed"]
        tbl = doc.add_table(rows=1, cols=len(cols))
        _add_table_borders(tbl)
        _header_row(tbl, cols)

        for idx, comp in enumerate(components):
            coa_req = "Yes" if comp.get("coa_required") else "No"
            _data_row(tbl, [
                comp.get("component", ""),
                f"{comp.get('part_number','')} {comp.get('revision','')}",
                comp.get("supplier", ""),
                comp.get("supplier_item_number", ""),
                comp.get("mock_lot", ""),
                "1",
                comp.get("material", ""),
                coa_req,
                f"Yes — CoA #{comp.get('mock_lot','')[:8].upper()}" if coa_req == "Yes" else "N/A",
            ], idx)
    else:
        doc.add_paragraph("No components defined in device profile.")

    # ── Section 3: In-Process Inspections ────────────────────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "Section 3 — In-Process Inspection Records")

    inspections = profile.get("in_process_inspections", [])
    if inspections:
        cols = ["Inspection Point", "Specification", "Method", "Actual Result",
                "Status", "Inspector", "Date"]
        tbl = doc.add_table(rows=1, cols=len(cols))
        _add_table_borders(tbl)
        _header_row(tbl, cols)

        for idx, insp in enumerate(inspections):
            _data_row(tbl, [
                insp.get("inspection_point", ""),
                insp.get("specification", ""),
                insp.get("method", ""),
                insp.get("actual_result", ""),
                insp.get("status", "PASS"),
                insp.get("inspector", ""),
                insp.get("date", ""),
            ], idx)
    else:
        doc.add_paragraph("No in-process inspections defined in device profile.")

    # ── Section 4: Environmental Conditions Log ───────────────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "Section 4 — Environmental Conditions Log")

    doc.add_paragraph(
        "Environmental monitoring data recorded during manufacturing "
        f"of Lot {lot}. Manufacturing area class: ISO Class 8 (Class 100,000) "
        "cleanroom unless otherwise noted."
    )

    mfg_date = profile.get("mock_manufacture_date", "2025-10-15")
    env_data = [
        [mfg_date, "07:00", "Assembly Area A1", "20.4°C", "46%",
         profile.get("personnel", {}).get("production_supervisor", "P. Supervisor")],
        [mfg_date, "10:00", "Assembly Area A1", "20.6°C", "47%",
         profile.get("personnel", {}).get("quality_engineer", "Q. Engineer")],
        [mfg_date, "13:00", "Assembly Area A1", "20.3°C", "45%",
         profile.get("personnel", {}).get("production_supervisor", "P. Supervisor")],
        [mfg_date, "16:00", "Assembly Area A1 + Clean Bench CB-01", "20.5°C", "46%",
         profile.get("personnel", {}).get("quality_engineer", "Q. Engineer")],
    ]
    if profile.get("is_sterile"):
        env_data.append(
            [mfg_date, "08:30", "Sterile Filling Area ISO 5", "19.8°C", "43%",
             profile.get("personnel", {}).get("production_supervisor", "P. Supervisor")]
        )

    cols = ["Date", "Time", "Area / Location", "Temperature (°C)", "Relative Humidity (%)",
            "Recorded By"]
    tbl = doc.add_table(rows=1, cols=len(cols))
    _add_table_borders(tbl)
    _header_row(tbl, cols)
    for idx, row_data in enumerate(env_data):
        _data_row(tbl, row_data, idx)

    # Temp/RH limits note
    doc.add_paragraph(
        "\nAcceptance Limits: Temperature 18–25°C | Relative Humidity ≤60%. "
        "All readings within limits for the manufacturing period. "
        "See Environmental Monitoring Log EM-" + mfg_date.replace("-", "") + " for full record."
    ).runs[0].font.size = Pt(9) if True else None

    # ── Section 5: Equipment Log ──────────────────────────────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "Section 5 — Equipment Identification and Calibration Status")

    equipment_seen = {}
    for step in profile.get("manufacturing_steps", []):
        eq_id   = step.get("equipment_id", "")
        eq_name = step.get("equipment", "")
        if eq_id and eq_id not in equipment_seen:
            equipment_seen[eq_id] = eq_name

    if equipment_seen:
        cols = ["Equipment ID", "Equipment Description", "Calibration / Qualification Status",
                "Calibration Due Date", "Certificate Number", "Used By (Operator)"]
        tbl = doc.add_table(rows=1, cols=len(cols))
        _add_table_borders(tbl)
        _header_row(tbl, cols)
        idx = 0
        for eq_id, eq_name in equipment_seen.items():
            cal_due = "2026-04-15"  # generic future date
            _data_row(tbl, [
                eq_id,
                eq_name,
                "Current — Calibrated",
                cal_due,
                f"CAL-{eq_id[-3:] if len(eq_id) >= 3 else '001'}-2025",
                profile.get("personnel", {}).get("quality_engineer", ""),
            ], idx)
            idx += 1

    # ── Section 6: Deviations and Nonconformances ─────────────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "Section 6 — Deviations and Nonconformances")

    # Add one minor deviation for training realism
    dev_p = doc.add_paragraph(
        f"One minor deviation was identified and dispositioned during the manufacture of Lot {lot}:"
    )

    dev_table_data = [
        ["Deviation Number", f"DEV-{lot}-001"],
        ["Date Identified", mfg_date],
        ["Operation / Step", profile.get("manufacturing_steps", [{}])[1].get("operation", "N/A") if len(profile.get("manufacturing_steps", [])) > 1 else "Final Inspection"],
        ["Description", "One unit identified with cosmetic blemish (scratch < 0.5 mm) on non-critical surface. Does not affect device function, sterility, or safety."],
        ["Disposition", "Use As Is — Approved per NCR Review Board"],
        ["Approved By", profile.get("personnel", {}).get("qa_manager", "")],
        ["NCR Reference", f"NCR-{lot}-001"],
        ["Impact on Product Quality", "None — cosmetic only. Device meets all functional acceptance criteria."],
    ]

    dev_tbl = doc.add_table(rows=len(dev_table_data), cols=2)
    _add_table_borders(dev_tbl)
    for ridx, (field, val) in enumerate(dev_table_data):
        row = dev_tbl.rows[ridx]
        row.cells[0].text = field
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = val
        if ridx % 2 == 0:
            _shade_cell(row.cells[0], "F2F2F2")
            _shade_cell(row.cells[1], "F2F2F2")

    doc.add_paragraph(
        f"\nAll other units in Lot {lot} ({profile.get('mock_batch_quantity','')} units total) "
        "met all acceptance criteria without deviation."
    )

    # ── Section 7: Labeling Review ────────────────────────────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "Section 7 — Labeling Review and Approval")

    label_items = [
        ("Device Name / Trade Name present and accurate",          "Yes"),
        ("Manufacturer name and address on label",                 "Yes"),
        ("Model / Part Number matches DMR",                        "Yes"),
        (f"Lot Number '{lot}' printed correctly",                  "Yes"),
        (f"Manufacture Date '{mfg_date}' on label",               "Yes"),
        (f"Expiration Date '{profile.get('mock_expiration_date','')}' on label",
         "Yes" if profile.get("mock_expiration_date") != "N/A" else "N/A — no expiry"),
        ("Quantity per package stated",                            "Yes"),
        ("Sterile / Non-Sterile designation correct",
         "Sterile (EO)" if profile.get("is_sterile") else "Non-Sterile — marked correctly"),
        ("UDI Device Identifier (DI) present",                     "Yes" if not profile.get("is_pharma") else "NDC present"),
        ("Regulatory (Rx Only / OTC) symbol correct",             "Yes" if profile.get("is_pharma") else "Rx Only symbol present" if profile.get("device_class") not in ["Class I"] else "No Rx symbol required"),
        ("Instructions for Use (IFU) reference included",          "Yes"),
        ("Label artwork matches released print file",              "Yes"),
        ("Barcode scanned and verified",                           "Yes"),
    ]

    cols = ["Labeling Check Item", "Status", "Reviewer Initials", "Date"]
    tbl = doc.add_table(rows=1, cols=len(cols))
    _add_table_borders(tbl)
    _header_row(tbl, cols)
    qa_initials = "".join(w[0].upper() for w in profile.get("personnel", {}).get("qa_manager", "Q M").split() if w)
    for idx, (item, status) in enumerate(label_items):
        _data_row(tbl, [item, status, qa_initials, mfg_date], idx)

    doc.add_paragraph(
        f"\nLabel review completed by {profile.get('personnel',{}).get('qa_manager','')} "
        f"on {mfg_date}. Label approved for release."
    )

    # ── Section 8: Sterilization (if applicable) ──────────────────────────────
    if profile.get("is_sterile"):
        doc.add_page_break()
        _add_section_heading(doc, "Section 8 — Sterilization Record")

        sterile_data = [
            ["Sterilization Method",         "Ethylene Oxide (EO)"],
            ["Sterilization Site",           "ContractSterile Partners LLC — 1800 Cleanroom Dr, San Jose, CA 95110"],
            ["Sterilization Work Order",     f"CSP-WO-{lot}"],
            ["Cycle Number",                 "Cycle EO-2025-1042"],
            ["Cycle Date",                   mfg_date],
            ["EO Concentration",             "600 mg/L (±50)"],
            ["Temperature",                  "54°C (±3°C)"],
            ["Relative Humidity",            "60% ±10%"],
            ["Exposure Time",                "180 minutes"],
            ["Degassing Period",             "14 days minimum at ambient conditions"],
            ["Bioburden Result (pre-sterile)", "≤ 50 CFU/device (acceptable limit: ≤ 100 CFU)"],
            ["SAL (Sterility Assurance Level)", "10⁻⁶ (demonstrated per ISO 11135)"],
            ["Biological Indicator (BI) Result", "All BIs killed — sterile"],
            ["Sterility Certificate Reference", f"CSP-CERT-{lot}"],
            ["EO Residual (ethylene oxide)", "≤ 4 ppm (limit: 4 ppm per ISO 10993-7)"],
        ]

        strl_tbl = doc.add_table(rows=len(sterile_data), cols=2)
        _add_table_borders(strl_tbl)
        for ridx, (field, val) in enumerate(sterile_data):
            row = strl_tbl.rows[ridx]
            row.cells[0].text = field
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = val
            if ridx % 2 == 0:
                _shade_cell(row.cells[0], "F2F2F2")
                _shade_cell(row.cells[1], "F2F2F2")

    # ── Section 9: Final Release Authorization ────────────────────────────────
    doc.add_page_break()
    _add_section_heading(doc, f"Section {'9' if profile.get('is_sterile') else '8'} — Final Release Authorization")

    qa_mgr   = profile.get("personnel", {}).get("qa_manager", "")
    prod_sup = profile.get("personnel", {}).get("production_supervisor", "")

    release_checklist = [
        ("All manufacturing steps completed and documented",                          "✓ Complete"),
        ("All in-process inspections passed",                                         "✓ Pass"),
        ("Bill of Materials — lot traceability complete",                             "✓ Complete"),
        ("Certificate(s) of Analysis reviewed and on file",                           "✓ On File"),
        ("Environmental monitoring within limits",                                    "✓ Pass"),
        ("All equipment in-calibration at time of manufacture",                       "✓ Verified"),
        ("Label review completed and approved",                                       "✓ Approved"),
        ("Sterilization certificate on file" if profile.get("is_sterile") else "Sterility — N/A",
         "✓ On File" if profile.get("is_sterile") else "N/A"),
        ("No open critical deviations",                                               "✓ No critical deviations"),
        ("Deviations dispositioned and closed (minor deviation documented above)",     "✓ Dispositioned"),
        ("Nonconformance review completed",                                           "✓ Complete"),
        ("All required records present in DHR",                                       "✓ Complete"),
    ]

    cols = ["Release Check Item", "Status"]
    tbl = doc.add_table(rows=1, cols=len(cols))
    _add_table_borders(tbl)
    _header_row(tbl, cols)
    for idx, (item, status) in enumerate(release_checklist):
        _data_row(tbl, [item, status], idx)

    doc.add_paragraph()

    release_stmt = doc.add_paragraph()
    r = release_stmt.add_run(
        f"RELEASE STATEMENT: Based on review of all manufacturing records contained in this "
        f"Device History Record, Lot {lot} of {profile.get('device_name','')} "
        f"({profile.get('trade_name','')}, Model {profile.get('model_number','')}) consisting of "
        f"{profile.get('mock_batch_quantity','')} units manufactured on {mfg_date} "
        f"is hereby RELEASED for distribution. All acceptance criteria have been met. "
        f"Any noted deviations have been reviewed and dispositioned as acceptable."
    )
    r.bold = True

    doc.add_paragraph()

    sig_data = [
        ["QA Manager:", qa_mgr, "Signature: ___________________", f"Date: {mfg_date}"],
        ["Production Supervisor:", prod_sup, "Signature: ___________________", f"Date: {mfg_date}"],
    ]
    sig_tbl = doc.add_table(rows=len(sig_data) + 1, cols=4)
    _add_table_borders(sig_tbl)
    for ci, hdr in enumerate(["Role", "Printed Name", "Signature", "Date"]):
        sig_tbl.rows[0].cells[ci].text = hdr
        sig_tbl.rows[0].cells[ci].paragraphs[0].runs[0].bold = True
        _shade_cell(sig_tbl.rows[0].cells[ci])
    for ri, row_vals in enumerate(sig_data, 1):
        for ci, val in enumerate(row_vals):
            sig_tbl.rows[ri].cells[ci].text = val

    # ── Disclaimer footer ─────────────────────────────────────────────────────
    _add_disclaimer_footer(doc)

    doc.save(str(out_path))
    log.info(f"DHR DOCX saved: {out_path.name}")


# ════════════════════════════════════════════════════════════════════════════
# MAIN SIMULATOR CLASS
# ════════════════════════════════════════════════════════════════════════════

class QMSSimulator:
    """Orchestrates the full QMS document bundle generation pipeline."""

    def __init__(self):
        if not ANTHROPIC_API_KEY:
            raise RuntimeError("ANTHROPIC_API_KEY not set. Check Athena/voice/.env")
        self.client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    # ── Internal helpers ───────────────────────────────────────────────────

    def _call(self, prompt: str, model: str = MODEL_DOCUMENT,
              max_tokens: int = 4000, system: str = None) -> str:
        """Single Claude API call; returns stripped text."""
        resp = self.client.messages.create(
            model=model,
            max_tokens=max_tokens,
            system=system or QMS_SYSTEM,
            messages=[{"role": "user", "content": prompt}],
        )
        text = resp.content[0].text.strip()
        if _MEM:
            try:
                _MEM.log_api_call("qms_simulator", model,
                                  resp.usage.input_tokens, resp.usage.output_tokens,
                                  purpose="qms_doc_gen")
            except Exception:
                pass
        return text

    def _parse_json(self, text: str) -> dict:
        """Extract and parse JSON from a Claude response."""
        if not text or not text.strip():
            raise json.JSONDecodeError("Empty response from Claude", "", 0)

        # Strategy 1: find the outermost { ... } block and parse it directly
        start = text.find("{")
        end   = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            candidate = text[start:end + 1]
            try:
                return json.loads(candidate)
            except json.JSONDecodeError:
                pass

        # Strategy 2: strip markdown code fences and try again
        cleaned = re.sub(r"```(?:json)?", "", text)
        cleaned = cleaned.strip()
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            pass

        log.error(f"JSON parse failed. Raw response (first 800 chars):\n{text[:800]}")
        raise json.JSONDecodeError("Could not extract valid JSON", text, 0)

    def _format_prompt(self, template: str, profile: dict) -> str:
        """Inject profile fields into a prompt template."""
        p = profile
        pers = p.get("personnel", {})
        return template.format(
            profile_json     = json.dumps(p, indent=2),
            device_name      = p.get("device_name", "Device"),
            model_number     = p.get("model_number", "MDL-001"),
            manufacture_date = p.get("mock_manufacture_date", "2025-10-15"),
            lot_number       = p.get("mock_lot_number", "LOT-001"),
            qa_manager       = pers.get("qa_manager", "QA Manager"),
            quality_engineer = pers.get("quality_engineer", "Quality Engineer"),
        )

    # ── Phase 1: Device Profile ────────────────────────────────────────────

    def generate_device_profile(self, description: str) -> dict:
        """
        Generate the master device profile JSON that drives all other documents.
        This is the most important call — everything else references this data.
        """
        log.info(f"Generating device profile for: {description!r}")
        prompt = PROFILE_PROMPT.format(description=description)
        raw    = self._call(prompt, model=MODEL_PROFILE, max_tokens=8000,
                            system=PROFILE_SYSTEM)
        try:
            profile = self._parse_json(raw)
            log.info(
                f"Profile generated: {profile.get('trade_name')} / "
                f"{profile.get('device_class')} / Lot {profile.get('mock_lot_number')}"
            )
            return profile
        except json.JSONDecodeError as e:
            log.error(f"Failed to parse device profile JSON: {e}")
            log.debug(f"Raw response (first 500 chars): {raw[:500]}")
            raise

    # ── Phase 2: Document generation ──────────────────────────────────────

    def generate_dmr_index(self, profile: dict) -> str:
        log.info("Generating DMR Index...")
        prompt = self._format_prompt(DMR_PROMPT, profile)
        return self._call(prompt) + DISCLAIMER_MD

    def generate_risk_summary(self, profile: dict) -> str:
        log.info("Generating Risk Management File Summary...")
        prompt = self._format_prompt(RISK_PROMPT, profile)
        return self._call(prompt) + DISCLAIMER_MD

    def generate_capa_records(self, profile: dict) -> str:
        log.info("Generating CAPA Records...")
        prompt = self._format_prompt(CAPA_PROMPT, profile)
        return self._call(prompt, max_tokens=5000) + DISCLAIMER_MD

    def generate_sop_list(self, profile: dict) -> str:
        log.info("Generating SOP Reference List...")
        prompt = self._format_prompt(SOP_PROMPT, profile)
        return self._call(prompt) + DISCLAIMER_MD

    def generate_training_records(self, profile: dict) -> str:
        log.info("Generating Training Records...")
        prompt = self._format_prompt(TRAINING_PROMPT, profile)
        return self._call(prompt) + DISCLAIMER_MD

    def generate_supplier_qual(self, profile: dict) -> str:
        log.info("Generating Supplier Qualification Records...")
        prompt = self._format_prompt(SUPPLIER_PROMPT, profile)
        return self._call(prompt) + DISCLAIMER_MD

    def generate_traceability(self, profile: dict) -> str:
        log.info("Generating Design & Risk Control Traceability Matrix...")
        prompt = self._format_prompt(TRACEABILITY_PROMPT, profile)
        return self._call(prompt, max_tokens=5000) + DISCLAIMER_MD

    def generate_risk_analysis(self, profile: dict) -> str:
        log.info("Generating Risk Analysis (Hazard Analysis + FMEA)...")
        prompt = self._format_prompt(RISK_ANALYSIS_PROMPT, profile)
        return self._call(prompt, max_tokens=5000) + DISCLAIMER_MD

    def generate_ncmr_trending(self, profile: dict) -> str:
        log.info("Generating NCMR / Complaint Trending Report...")
        prompt = self._format_prompt(NCMR_TRENDING_PROMPT, profile)
        return self._call(prompt, max_tokens=5000) + DISCLAIMER_MD

    # ── Phase 3: File I/O ──────────────────────────────────────────────────

    def save_markdown(self, content: str, path: Path):
        path.write_text(content, encoding="utf-8")
        log.info(f"  Saved: {path.name}")

    def build_bundle_index(self, profile: dict, out_dir: Path, files: dict) -> str:
        p = profile
        pers = p.get("personnel", {})
        lines = [
            "SIMULATED RECORDS — MOCK AUDIT USE ONLY",
            "",
            f"# QMS Document Bundle Index",
            f"**Device:** {p.get('device_name','')} — {p.get('trade_name','')} (Model {p.get('model_number','')})",
            f"**Manufacturer (fictional):** {p.get('manufacturer_name','')}",
            f"**Lot / Batch:** {p.get('mock_lot_number','')} | **Manufacture Date:** {p.get('mock_manufacture_date','')}",
            f"**Device Class:** {p.get('device_class','')} | **Regulatory Pathway:** {p.get('regulatory_pathway','')}",
            f"**Clearance Number (fictional):** {p.get('clearance_number','')}",
            f"**Bundle Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            f"**Generated by:** Latitude MedTech QMS Simulator Agent",
            "",
            "---",
            "",
            "## Document Inventory",
            "",
            "| # | File | Description | Format |",
            "|---|---|---|---|",
        ]
        doc_list = [
            ("00_DEVICE_PROFILE.json",         "Device Classification + Master Metadata",         "JSON"),
            ("DHR_batch.docx",                 "Device History Record (Primary Deliverable)",      "DOCX"),
            ("DMR_index.md",                   "Device Master Record Index",                       "Markdown"),
            ("risk_management_summary.md",     "ISO 14971 Risk Management File Summary",           "Markdown"),
            ("risk_analysis_fmea.md",          "Hazard Analysis + Design & Process FMEA",         "Markdown"),
            ("design_risk_traceability.md",    "Design Input/Output + Risk Control Traceability",  "Markdown"),
            ("CAPA_simulation.md",             "CAPA Records — 2–3 Mock Scenarios",                "Markdown"),
            ("NCMR_complaint_trending.md",     "NCMR Log + Complaint Trending + MDR Assessment",   "Markdown"),
            ("sop_reference_list.md",          "Applicable SOP Matrix",                            "Markdown"),
            ("training_records.md",            "Personnel Training Qualification Summary",         "Markdown"),
            ("supplier_qualification.md",      "Approved Supplier List + Qualification Data",      "Markdown"),
        ]
        for i, (fname, desc, fmt) in enumerate(doc_list, 1):
            lines.append(f"| {i} | `{fname}` | {desc} | {fmt} |")

        lines += [
            "",
            "---",
            "",
            "## Key Personnel (Fictional)",
            "",
            f"| Role | Name |",
            f"|---|---|",
            f"| QA Manager | {pers.get('qa_manager','')} |",
            f"| Production Supervisor | {pers.get('production_supervisor','')} |",
            f"| Quality Engineer | {pers.get('quality_engineer','')} |",
            f"| Regulatory Affairs | {pers.get('regulatory_affairs','')} |",
            f"| Design Engineer | {pers.get('design_engineer','N/A')} |",
            "",
            "---",
            "",
            "## Mock Audit Use Notes",
            "",
            "This document bundle is designed for mock ISO 13485:2016 audit training.",
            "Suggested mock audit activities:",
            "",
            "1.  **DHR Walkthrough** — Trace Lot through manufacturing traveler, verify BOM traceability",
            "2.  **CAPA Deep-Dive** — Audit CAPA records: is root cause adequate? Is verification measurable?",
            "3.  **Supplier Audit** — Challenge the Approved Supplier List: are critical suppliers qualified?",
            "4.  **DMR vs DHR** — Confirm the DHR references the correct DMR document revisions",
            "5.  **Risk Management Summary** — Review risk summary against DHR deviations and NCMR data",
            "6.  **Risk Analysis / FMEA** — Trace high-RPN FMEA items to risk controls; verify residual risk",
            "7.  **Traceability Challenge** — Pick a risk control: trace it from Hazard ID → DMR doc → DHR verification",
            "8.  **NCMR / Complaint Trending** — Are signal thresholds met? Were MDRs filed on time?",
            "9.  **Training Audit** — Verify all operators in the manufacturing traveler have current training records",
            "10. **Data Analysis (§8.4)** — Using the NCMR trending KPI table, would you open a CAPA?",
            "",
        ]
        lines.append(DISCLAIMER_MD)
        return "\n".join(lines)

    # ── Main orchestration ─────────────────────────────────────────────────

    def run(self, description, out_dir=None, profile_only=False):
        """
        Full pipeline: profile → DHR DOCX → 9 Markdown documents → bundle index.
        Returns the output directory path.
        """
        ts    = datetime.now().strftime("%Y%m%d_%H%M%S")
        slug  = re.sub(r"[^a-z0-9]+", "_", description.lower())[:30].strip("_")
        run_dir = (out_dir or QMS_DIR) / f"{ts}_{slug}"
        run_dir.mkdir(parents=True, exist_ok=True)

        log.info(f"{'='*60}")
        log.info(f"QMS SIMULATOR — Starting run")
        log.info(f"Device description: {description!r}")
        log.info(f"Output directory:   {run_dir}")
        log.info(f"{'='*60}")

        # ── Step 1: Device Profile ────────────────────────────────────────
        profile    = self.generate_device_profile(description)
        lot        = profile.get("mock_lot_number", "LOT-001")
        trade_name = profile.get("trade_name", "")

        profile_path = run_dir / "00_DEVICE_PROFILE.json"
        profile_path.write_text(json.dumps(profile, indent=2), encoding="utf-8")
        log.info(f"  Saved: {profile_path.name}")

        if profile_only:
            log.info("--profile-only flag set. Stopping after device profile.")
            return run_dir

        # ── Step 2: DHR (Word document) ───────────────────────────────────
        dhr_path = run_dir / f"DHR_batch.docx"
        if _docx_available():
            try:
                build_dhr_docx(profile, dhr_path)
            except Exception as e:
                log.error(f"DHR DOCX build failed: {e}. Saving profile JSON only for DHR.")
        else:
            log.warning(
                "python-docx not available. Run with Athena/voice/venv to generate DHR DOCX. "
                "Use: .\\Athena\\voice\\venv\\Scripts\\python.exe qms_simulator_agent.py ..."
            )

        # ── Steps 3–11: Markdown documents ────────────────────────────────
        docs = {
            "DMR_index.md":                  self.generate_dmr_index(profile),
            "risk_management_summary.md":    self.generate_risk_summary(profile),
            "risk_analysis_fmea.md":         self.generate_risk_analysis(profile),
            "design_risk_traceability.md":   self.generate_traceability(profile),
            "CAPA_simulation.md":            self.generate_capa_records(profile),
            "NCMR_complaint_trending.md":    self.generate_ncmr_trending(profile),
            "sop_reference_list.md":         self.generate_sop_list(profile),
            "training_records.md":           self.generate_training_records(profile),
            "supplier_qualification.md":     self.generate_supplier_qual(profile),
        }

        for fname, content in docs.items():
            self.save_markdown(content, run_dir / fname)

        # ── Step 9: Bundle Index ──────────────────────────────────────────
        index_content = self.build_bundle_index(profile, run_dir, docs)
        self.save_markdown(index_content, run_dir / "QMS_BUNDLE_INDEX.md")

        # ── Memory log ────────────────────────────────────────────────────
        if _MEM:
            try:
                _MEM.log_event(
                    "qms_simulator", "bundle_generated",
                    subject=f"{trade_name} / {lot}",
                    metadata={"device": description, "path": str(run_dir), "docs": len(docs) + 2}
                )
            except Exception:
                pass

        log.info(f"{'='*60}")
        log.info(f"QMS BUNDLE COMPLETE — {len(docs) + 2} files (11 documents)")
        log.info(f"Output: {run_dir}")
        log.info(f"{'='*60}")

        return run_dir


# ════════════════════════════════════════════════════════════════════════════
# CLI ENTRY POINT
# ════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        prog="qms_simulator_agent.py",
        description=(
            "Latitude MedTech QMS Simulator — Generate a complete ISO 13485:2016 "
            "QMS document bundle (DHR DOCX + 9 supporting documents) for mock audit training."
        ),
        epilog=(
            "DISCLAIMER: All generated content is fictional. For educational / mock audit "
            "training use only. Not regulatory or legal advice."
        )
    )
    parser.add_argument(
        "description",
        help='Device or compound description (e.g., "Insulin Delivery Pen")',
    )
    parser.add_argument(
        "--out", metavar="DIR",
        help=f"Output directory (default: {QMS_DIR})",
    )
    parser.add_argument(
        "--profile-only", action="store_true",
        help="Stop after generating the device profile JSON (no document generation)",
    )
    args = parser.parse_args()

    out_dir = Path(args.out) if args.out else None

    sim = QMSSimulator()
    result_dir = sim.run(
        description  = args.description,
        out_dir      = out_dir,
        profile_only = args.profile_only,
    )

    print(f"\n{'='*60}")
    print(f"  QMS Bundle ready for mock audit use.")
    print(f"  Output directory: {result_dir}")
    print(f"{'='*60}\n")

    # List files
    for f in sorted(result_dir.iterdir()):
        size_kb = f.stat().st_size // 1024 if f.exists() else 0
        print(f"  {f.name:<45} {size_kb:>4} KB")

    print(f"\n  DISCLAIMER: All content is fictional. Mock audit training use only.")
    print(f"  Label: Alpha — Steve Review Required.\n")


if __name__ == "__main__":
    main()
