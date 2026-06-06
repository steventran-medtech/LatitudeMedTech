"""
Latitude MedTech — Regulatory Strategy Agent
=============================================
Core consulting product: gap assessment + remediation roadmap against
21 CFR Part 820, ISO 13485:2016, EU MDR 2017/745, and/or MDSAP.

CLI usage (invoked via server run_agent):
    python regulatory_strategy_agent.py \\
        --device-type "Class II cardiac monitor" \\
        --classification "Class II" \\
        --markets US EU \\
        --qms-state "ISO 13485 certified, no EU MDR technical file" \\
        [--client-id <id>]

Output: branded gap assessment .docx submitted to review queue.
"""

import os
import sys
import argparse
import logging
import re
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

from pathconfig import ATHENA_ROOT, ENV_FILE, AGENTS_DIR, LOGS_DIR
load_dotenv(ENV_FILE)
sys.path.insert(0, str(AGENTS_DIR))

from memory import Memory
from agent_base import AgentBase

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s  %(levelname)s  %(message)s',
    handlers=[
        logging.FileHandler(LOGS_DIR / 'regulatory_strategy_agent.log'),
        logging.StreamHandler(sys.stdout),
    ]
)
log = logging.getLogger('regulatory_strategy_agent')

DOCS_DIR = ATHENA_ROOT / 'documents' / 'regulatory'
DOCS_DIR.mkdir(parents=True, exist_ok=True)

FRAMEWORKS = {
    'US':    ['21 CFR Part 820 (QSR)', '21 CFR Part 803 (MDR Reporting)', 'FDA 510(k)/PMA pathway'],
    'EU':    ['EU MDR 2017/745', 'MDCG guidance', 'ISO 13485:2016'],
    'CA':    ['MDSAP', 'Health Canada SOR/98-282'],
    'AU':    ['MDSAP', 'TGA ARTG requirements'],
    'MDSAP': ['MDSAP audit approach', 'ISO 13485:2016'],
    'ISO':   ['ISO 13485:2016', 'ISO 14971:2019'],
}

REGULATORY_SYSTEM = """You are a senior regulatory affairs consultant at Latitude MedTech LLC.
You have 20+ years of experience with FDA, EU MDR, ISO 13485, and MDSAP.

Output style:
- Big 4 management consulting quality (McKinsey/PwC register)
- Outcome-oriented titles; Situation → Approach → Impact
- Every finding must cite the specific clause or regulation (e.g., "ISO 13485 §7.3.2")
- Remediation items must be concrete and time-bound
- Use tables in Markdown format for the gap matrix
- Plain direct voice; no motivational filler; no em-dash tics
- Label all outputs: Alpha — Steve Review Required
- Append standard disclaimer
"""

GAP_ASSESSMENT_PROMPT = """Conduct a regulatory gap assessment for the following device and current QMS state.

DEVICE PROFILE:
- Device type: {device_type}
- Risk classification: {classification}
- Target markets: {markets}
- Current QMS state: {qms_state}
- Additional context: {context}

APPLICABLE FRAMEWORKS:
{framework_list}

KNOWLEDGE BASE CONTEXT (use for grounding — cite regulatory sources):
{kb_context}

Generate a complete Regulatory Gap Assessment Report with these sections:

## Executive Summary
3–4 sentences: device profile, markets targeted, overall gap severity (Low/Medium/High/Critical), and highest-priority action.

## Regulatory Pathway Analysis
For each target market: applicable regulatory pathway, classification basis, and estimated timeline to market entry.

## Gap Assessment Matrix
Use a Markdown table with columns: | Requirement | Clause | Current State | Gap | Priority |
Priority = Critical / High / Medium / Low.
Cover at minimum: Design Controls, Risk Management, CAPA, Document Control, Supplier Management, Post-Market Surveillance, Labeling/UDI.

## Remediation Roadmap
Numbered list of remediation actions, grouped by:
### Phase 1 — Immediate (0–30 days): Critical gaps blocking submission
### Phase 2 — Near-term (30–90 days): High-priority gaps affecting audit readiness
### Phase 3 — Long-term (90–180 days): Medium/Low gaps for full compliance maturity

Each action: specific deliverable name + responsible party (e.g., "RA Lead") + estimated effort.

## Resource Estimate
Table: | Activity | Estimated Hours | Latitude Support Level |
Total estimated project hours and engagement investment range.

## Next Steps
3 specific actions for the client to take in the next 2 weeks.

DISCLAIMER: This gap assessment is produced by Latitude MedTech LLC for planning purposes only.
Not regulatory or legal advice. Findings require validation against current device-specific documentation.
Alpha — Steve Review Required.
"""


def build_regulatory_docx(title: str, content: str, device_type: str) -> Path:
    """Build a branded regulatory gap assessment .docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    LM_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
    LM_SLATE = RGBColor(0x2C, 0x3E, 0x50)
    LM_BLUE  = RGBColor(0x5B, 0x7F, 0xA6)
    LM_MUTED = RGBColor(0x8A, 0x86, 0x80)
    LM_GOLD  = RGBColor(0xC4, 0x92, 0x2A)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Header
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("LATITUDE MEDTECH LLC")
    run.font.name = "Calibri"; run.font.size = Pt(7.5)
    run.font.color.rgb = LM_BLUE

    # Footer
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(
        f"latitudemedtech.com  |  Regulatory Gap Assessment  |  "
        f"{datetime.now().strftime('%B %d, %Y')}  |  Confidential"
    )
    run.font.name = "Calibri"; run.font.size = Pt(7.5)
    run.font.color.rgb = LM_MUTED

    # Title block
    p = doc.add_paragraph()
    r = p.add_run("REGULATORY GAP ASSESSMENT")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = LM_BLUE; r.font.bold = True; r.font.all_caps = True

    p2 = doc.add_paragraph()
    r2 = p2.add_run(title)
    r2.font.name = "Calibri"; r2.font.size = Pt(16)
    r2.font.color.rgb = LM_BLACK; r2.font.bold = True
    p2.paragraph_format.space_after = Pt(4)

    p3 = doc.add_paragraph()
    r3 = p3.add_run(
        f"Device: {device_type}  |  Latitude MedTech LLC  |  "
        f"{datetime.now().strftime('%B %Y')}"
    )
    r3.font.name = "Calibri"; r3.font.size = Pt(10)
    r3.font.color.rgb = LM_MUTED
    p3.paragraph_format.space_after = Pt(14)

    # Parse and render markdown
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            doc.add_paragraph()
            i += 1
            continue

        # Table detection
        if '|' in line and i + 1 < len(lines) and re.match(r'^\|[-\s|]+\|', lines[i + 1]):
            # Collect table rows
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                if not re.match(r'^\|[-\s|]+\|', lines[i]):
                    table_lines.append(lines[i])
                i += 1
            if table_lines:
                headers = [c.strip() for c in table_lines[0].strip('|').split('|')]
                tbl = doc.add_table(rows=1, cols=len(headers))
                tbl.style = 'Table Grid'
                hrow = tbl.rows[0]
                for j, h in enumerate(headers):
                    cell = hrow.cells[j]
                    cell.text = h
                    run = cell.paragraphs[0].runs[0]
                    run.font.bold = True
                    run.font.name = "Calibri"; run.font.size = Pt(9)
                    run.font.color.rgb = LM_BLUE
                for row_line in table_lines[1:]:
                    cells = [c.strip() for c in row_line.strip('|').split('|')]
                    row = tbl.add_row()
                    for j, cell_text in enumerate(cells):
                        if j < len(row.cells):
                            row.cells[j].text = cell_text
                            run = row.cells[j].paragraphs[0].runs[0] if row.cells[j].paragraphs[0].runs else row.cells[j].paragraphs[0].add_run(cell_text)
                            run.font.name = "Calibri"; run.font.size = Pt(9)
                            run.font.color.rgb = LM_BLACK
                doc.add_paragraph()
            continue

        if line.startswith('### '):
            p = doc.add_paragraph()
            r = p.add_run(line[4:])
            r.font.name = "Calibri"; r.font.size = Pt(11)
            r.font.color.rgb = LM_SLATE; r.font.bold = True
            p.paragraph_format.space_before = Pt(8)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            r = p.add_run(line[3:].upper())
            r.font.name = "Calibri"; r.font.size = Pt(11)
            r.font.color.rgb = LM_BLUE; r.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
        elif line.startswith('# '):
            p = doc.add_paragraph()
            r = p.add_run(line[2:])
            r.font.name = "Calibri"; r.font.size = Pt(14)
            r.font.color.rgb = LM_BLACK; r.font.bold = True
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line[2:])
            r = p.add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(10)
            r.font.color.rgb = LM_BLACK
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\. ', '', line)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            r = p.add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(10)
            r.font.color.rgb = LM_BLACK
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line[2:])
            r = p.add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(10)
            r.font.color.rgb = LM_SLATE; r.font.italic = True
        else:
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(10)
            r.font.color.rgb = LM_BLACK
            p.paragraph_format.space_after = Pt(4)

        i += 1

    slug = re.sub(r'[^\w]', '_', device_type.lower())[:30]
    ts   = datetime.now().strftime('%Y%m%d_%H%M%S')
    out  = DOCS_DIR / f"gap_assessment_{slug}_{ts}.docx"
    doc.save(str(out))
    return out


def run_assessment(device_type: str, classification: str, markets: list,
                   qms_state: str, context: str = "", client_id: int = None):
    mem  = Memory()
    base = AgentBase("regulatory_strategy_agent")

    # Build framework list from target markets
    applicable = set()
    for market in markets:
        applicable.update(FRAMEWORKS.get(market.upper(), []))
    framework_list = "\n".join(f"- {f}" for f in sorted(applicable))

    # KB retrieval — search across regulatory + QMS topics
    query    = f"{device_type} {classification} {' '.join(markets)} gap assessment regulatory"
    kb_ctx   = base.central_kb_context(query, top_k=6, max_chars=3000)

    prompt = GAP_ASSESSMENT_PROMPT.format(
        device_type=device_type,
        classification=classification,
        markets=', '.join(markets),
        qms_state=qms_state,
        context=context or "N/A",
        framework_list=framework_list,
        kb_context=kb_ctx or "No pre-retrieved context available — rely on regulatory knowledge.",
    )

    log.info(f"Running gap assessment: {device_type} / {classification} / {markets}")

    client_obj = base._get_client()
    resp = client_obj.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        system=REGULATORY_SYSTEM,
        messages=[{"role": "user", "content": prompt}],
    )
    assessment_content = resp.content[0].text.strip()
    base.log_api("claude-sonnet-4-6", resp.usage.input_tokens,
                 resp.usage.output_tokens, purpose="regulatory_gap_assessment")

    title = f"Regulatory Gap Assessment — {device_type} — {', '.join(markets)}"
    doc_path = build_regulatory_docx(title, assessment_content, device_type)
    log.info(f"Gap assessment saved: {doc_path}")

    # Link to client engagement if provided
    if client_id:
        eng_id = mem.add_engagement(
            client_id=client_id,
            title=title,
            description=f"{classification} {device_type} — {', '.join(markets)} regulatory gap assessment",
            sow_path=str(doc_path),
            status='deliverable_ready',
        )

    review_id = mem.submit_for_review(
        agent="regulatory_strategy_agent",
        item_type="gap_assessment",
        title=title,
        file_path=str(doc_path),
    )
    mem.log_event(
        "regulatory_strategy_agent", "gap_assessment_generated",
        subject=device_type,
        metadata={
            "classification": classification,
            "markets": markets,
            "client_id": client_id,
            "review_id": review_id,
        }
    )
    log.info(f"Gap assessment queued for review (id={review_id})")
    print(f"ASSESSMENT_GENERATED: {doc_path.name}")


def main():
    parser = argparse.ArgumentParser(description="Latitude Regulatory Strategy Agent")
    parser.add_argument('--device-type',    required=True, help="Device description (e.g. 'Class II cardiac monitor')")
    parser.add_argument('--classification', required=True, help="Risk class (e.g. 'Class II', 'Class IIa')")
    parser.add_argument('--markets',        nargs='+', default=['US'], help="Target markets (US EU CA AU ISO MDSAP)")
    parser.add_argument('--qms-state',      required=True, help="Current QMS state description")
    parser.add_argument('--context',        default="", help="Additional engagement context")
    parser.add_argument('--client-id',      type=int, default=None, help="Link to client record")
    args = parser.parse_args()

    run_assessment(
        device_type=args.device_type,
        classification=args.classification,
        markets=args.markets,
        qms_state=args.qms_state,
        context=args.context,
        client_id=args.client_id,
    )


if __name__ == "__main__":
    main()
