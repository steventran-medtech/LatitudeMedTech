"""
Latitude MedTech — Statement of Work (SOW) Generation Agent
=============================================================
Generates a branded Statement of Work .docx from client intake data.

CLI usage (invoked via server run_agent):
    python sow_agent.py --client-id <id> [--engagement-id <id>]

The agent reads client + engagement data from the DB, generates SOW
content with Claude Sonnet, builds a branded .docx, saves it, and
submits to the review queue.
"""

import os
import sys
import argparse
import logging
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

from pathconfig import ATHENA_ROOT, ENV_FILE, AGENTS_DIR, LOGS_DIR
load_dotenv(ENV_FILE)
sys.path.insert(0, str(AGENTS_DIR))

from memory import Memory
from agent_base import AgentBase

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s  %(levelname)s  %(message)s',
    handlers=[
        logging.FileHandler(LOGS_DIR / 'sow_agent.log'),
        logging.StreamHandler(sys.stdout),
    ]
)
log = logging.getLogger('sow_agent')

DOCS_DIR = ATHENA_ROOT / 'documents' / 'sow'
DOCS_DIR.mkdir(parents=True, exist_ok=True)

PROGRAM_TIERS = {
    'career_prep':        {'label': 'Career Prep Program',        'price': 699,   'duration': '4 weeks'},
    'early_career':       {'label': 'Early Career QA/RA Program', 'price': 1499,  'duration': '8 weeks'},
    'mid_career':         {'label': 'Mid-Career Acceleration',    'price': 1899,  'duration': '10 weeks'},
    'career_transition':  {'label': 'Career Transition Program',  'price': 2299,  'duration': '12 weeks'},
    'consulting_gap':     {'label': 'Regulatory Gap Assessment',  'price': 4500,  'duration': '3–4 weeks'},
    'consulting_strategy':{'label': 'Regulatory Strategy Engagement', 'price': 8500, 'duration': '6–8 weeks'},
    'custom':             {'label': 'Custom Engagement',          'price': None,  'duration': 'TBD'},
}

SOW_SYSTEM = """You are a Big 4-quality management consulting proposal writer for Latitude MedTech LLC.
Generate a professional Statement of Work in structured Markdown.

House style rules:
- Outcome-oriented titles with quantified scope where possible
- Situation → Approach → Impact structure
- Plain, direct voice — no motivational filler, no em-dash tics
- Every section has a defensible claim or concrete deliverable
- All outputs labeled: Alpha — Steve Review Required
- Append disclaimer: "This Statement of Work is subject to review and approval by Latitude MedTech LLC before execution. Not regulatory or legal advice."
"""

SOW_PROMPT = """Generate a complete Statement of Work for the following client engagement.

CLIENT INFORMATION:
{client_block}

ENGAGEMENT:
{engagement_block}

PROGRAM TIER: {tier_label} ({tier_price})

Generate a Statement of Work with these sections:
1. Engagement Overview (2–3 sentences: client context + what we're solving)
2. Scope of Work (3–5 specific deliverables as a numbered list)
3. Engagement Approach (methodology — discovery → analysis → deliverable → review)
4. Timeline (phase breakdown for {duration})
5. Investment (program price, payment structure: 50% upfront / 50% on delivery)
6. Terms & Conditions (4–6 bullet points: IP ownership, confidentiality, revision rounds, cancellation)
7. Acceptance (signature block placeholders for Steven Tran and Client)

Use Markdown formatting. Be specific and concrete — no generic consulting filler.
All deliverables should be named artifacts (e.g. "Gap Assessment Report", "Remediation Roadmap").
"""


def build_sow_docx(title: str, content: str, client_name: str) -> Path:
    """Build a branded .docx SOW using python-docx directly."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re

    LM_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
    LM_SLATE = RGBColor(0x2C, 0x3E, 0x50)
    LM_BLUE  = RGBColor(0x5B, 0x7F, 0xA6)
    LM_MUTED = RGBColor(0x8A, 0x86, 0x80)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Header
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("LATITUDE MEDTECH LLC")
    run.font.name = "Calibri"; run.font.size = Pt(7.5)
    run.font.color.rgb = LM_BLUE

    # Footer
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(
        f"latitudemedtech.com  |  Statement of Work  |  {datetime.now().strftime('%B %d, %Y')}  |  "
        "Confidential — Not for Distribution"
    )
    run.font.name = "Calibri"; run.font.size = Pt(7.5)
    run.font.color.rgb = LM_MUTED

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("STATEMENT OF WORK")
    r.font.name = "Calibri"; r.font.size = Pt(9)
    r.font.color.rgb = LM_BLUE; r.font.bold = True; r.font.all_caps = True

    p2 = doc.add_paragraph()
    r2 = p2.add_run(title)
    r2.font.name = "Calibri"; r2.font.size = Pt(18)
    r2.font.color.rgb = LM_BLACK; r2.font.bold = True
    p2.paragraph_format.space_after = Pt(4)

    p3 = doc.add_paragraph()
    r3 = p3.add_run(f"Prepared for: {client_name}  |  Latitude MedTech LLC  |  {datetime.now().strftime('%B %Y')}")
    r3.font.name = "Calibri"; r3.font.size = Pt(10)
    r3.font.color.rgb = LM_MUTED
    p3.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()  # spacer

    # Parse and render markdown content
    lines = content.split('\n')
    for line in lines:
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith('## '):
            p = doc.add_paragraph()
            r = p.add_run(line[3:].upper())
            r.font.name = "Calibri"; r.font.size = Pt(11)
            r.font.color.rgb = LM_BLUE; r.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
        elif line.startswith('# '):
            p = doc.add_paragraph()
            r = p.add_run(line[2:])
            r.font.name = "Calibri"; r.font.size = Pt(13)
            r.font.color.rgb = LM_SLATE; r.font.bold = True
            p.paragraph_format.space_before = Pt(10)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            r = p.add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(10)
            r.font.color.rgb = LM_BLACK
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\. ', '', line)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            r = p.add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(10)
            r.font.color.rgb = LM_BLACK
        else:
            # Strip inline markdown
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(10)
            r.font.color.rgb = LM_BLACK
            p.paragraph_format.space_after = Pt(4)

    # Disclaimer
    doc.add_paragraph()
    pd = doc.add_paragraph()
    rd = pd.add_run(
        "DISCLAIMER: This Statement of Work is subject to review and approval by "
        "Latitude MedTech LLC before execution. All content is produced for planning "
        "purposes only and does not constitute regulatory or legal advice. "
        "Alpha — Steve Review Required."
    )
    rd.font.name = "Calibri"; rd.font.size = Pt(8)
    rd.font.color.rgb = LM_MUTED; rd.font.italic = True

    slug = re.sub(r'[^\w]', '_', client_name.lower())[:30]
    ts   = datetime.now().strftime('%Y%m%d_%H%M%S')
    out  = DOCS_DIR / f"sow_{slug}_{ts}.docx"
    doc.save(str(out))
    return out


def generate_sow(client_id: int, engagement_id: int = None):
    mem  = Memory()
    base = AgentBase("sow_agent")

    client = mem.get_client(client_id)
    if not client:
        log.error(f"Client {client_id} not found")
        sys.exit(1)

    engagement = None
    if engagement_id:
        engagements = mem.get_engagements(client_id)
        engagement  = next((e for e in engagements if e['id'] == engagement_id), None)

    tier_key  = client.get('program_tier') or 'custom'
    tier_info = PROGRAM_TIERS.get(tier_key, PROGRAM_TIERS['custom'])

    client_block = "\n".join(filter(None, [
        f"Name: {client['name']}",
        f"Organization: {client.get('org', 'N/A')}",
        f"Role: {client.get('role', 'N/A')}",
        f"Regulatory challenge: {client.get('regulatory_challenge', 'N/A')}",
        f"Timeline: {client.get('timeline', 'N/A')}",
        f"Budget range: {client.get('budget_range', 'N/A')}",
        f"Notes: {client.get('notes', '')}" if client.get('notes') else None,
    ]))

    engagement_block = ""
    if engagement:
        engagement_block = "\n".join(filter(None, [
            f"Title: {engagement['title']}",
            f"Description: {engagement.get('description', '')}",
            f"Status: {engagement.get('status', 'scoping')}",
        ]))
    else:
        engagement_block = f"Scope: {client.get('regulatory_challenge', 'To be determined')}"

    price_str = f"${tier_info['price']:,}" if tier_info['price'] else "Custom pricing"

    prompt = SOW_PROMPT.format(
        client_block=client_block,
        engagement_block=engagement_block,
        tier_label=tier_info['label'],
        tier_price=price_str,
        duration=tier_info['duration'],
    )

    kb_ctx = base.kb_context(client.get('regulatory_challenge', 'consulting services'), top_k=3)
    system = SOW_SYSTEM + (f"\n\n{kb_ctx}" if kb_ctx else "")

    log.info(f"Generating SOW for client: {client['name']} (tier: {tier_key})")

    client_obj = base._get_client()
    resp = client_obj.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2500,
        system=system,
        messages=[{"role": "user", "content": prompt}],
    )
    sow_content = resp.content[0].text.strip()
    base.log_api("claude-sonnet-4-6", resp.usage.input_tokens,
                 resp.usage.output_tokens, purpose="sow_generation")

    title = f"Statement of Work — {client['name']} — {tier_info['label']}"
    doc_path = build_sow_docx(title, sow_content, client['name'])
    log.info(f"SOW saved: {doc_path}")

    # Update engagement with SOW path
    if engagement_id:
        mem.update_engagement(engagement_id, sow_path=str(doc_path), status='sow_generated')
    else:
        # Create a new engagement record
        eng_id = mem.add_engagement(
            client_id=client_id,
            title=title,
            description=client.get('regulatory_challenge', ''),
            sow_path=str(doc_path),
            status='sow_generated',
            value_usd=tier_info['price'],
        )

    review_id = mem.submit_for_review(
        agent="sow_agent",
        item_type="sow",
        title=title,
        file_path=str(doc_path),
    )
    mem.log_event("sow_agent", "sow_generated", subject=client['name'],
                  metadata={"client_id": client_id, "tier": tier_key, "review_id": review_id})
    log.info(f"SOW queued for review (id={review_id})")
    print(f"SOW_GENERATED: {doc_path.name}")


def main():
    parser = argparse.ArgumentParser(description="Latitude SOW Agent")
    parser.add_argument('--client-id',    type=int, required=True)
    parser.add_argument('--engagement-id',type=int, default=None)
    args = parser.parse_args()
    generate_sow(args.client_id, args.engagement_id)


if __name__ == "__main__":
    main()
